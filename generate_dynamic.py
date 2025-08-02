import pandas as pd
from docx import Document
import os
import logging
from tkinter import Tk, filedialog
from docx2pdf import convert
from datetime import datetime, date
from data_folder.send_email import send_email_with_attachment
from data_folder.helper_file import get_body, is_valid_email_syntax

# create a base folder for log and outputs
output_folder = f'receipts_{datetime.now().strftime('%Y-%m-%d_%H%M%S')}'
os.makedirs(output_folder, exist_ok=True)  # Create folder if it doesn't exist

log_file = os.path.join(output_folder, f"log_{datetime.now().strftime('%Y-%m-%d_%H%M%S')}")
# Configure logging
logging.basicConfig(
    filename=log_file,
    filemode='a',
    format='%(asctime)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logging.info('Script started.')
# --- Select Excel File ---
print('Prompting user to upload excel file for data.')
Tk().withdraw()
excel_file = filedialog.askopenfilename(
    title="Select Excel File",
    filetypes=[("Excel files", "*.xlsx *.xls")]
)
logging.info(f'Excel file selected: {excel_file}')
if not excel_file:
    print("❌ No Excel file selected.")
    logging.error("❌ No Excel file selected.")
    exit()


# --- Select Word Template ---
print('Prompting user to upload word file for template.')
template_file = filedialog.askopenfilename(
    title="Select Word Template (.docx)",
    filetypes=[("Word Documents", "*.docx *.doc")]
)
logging.info(f'Word template selected: {template_file}')
if not template_file:
    print("❌ No Word template selected.")
    logging.error("❌ No Word template selected.")
    exit()

initial_email = input("Receipts will be generated. Do you wish to initiate email delivery? (Y/N): ").lower() == 'y'
logging.info(f'Initiation for emails: {initial_email}')

body_file_path = get_body()
logging.info(f'Email body file path: {body_file_path}')


# --- Load Data ---
logging.info(f'Loading data from {excel_file}')
df = pd.read_excel(excel_file)
headers = df.columns.tolist()
logging.info(f'Headers found: {headers}')

# --- Output Directory ---

output_dir_word = os.path.join(output_folder,"generated_docs")
os.makedirs(output_dir_word, exist_ok=True)
logging.info(f'Word output directory: {output_dir_word}')

output_dir_pdf = os.path.join(output_folder,"generated_pdf")
os.makedirs(output_dir_pdf, exist_ok=True)
logging.info(f'PDF output directory: {output_dir_pdf}')

def format_value(value):
    try:
        if value == 'nan' or pd.isna(value):
            return ""
        elif isinstance(value, date):
            val = datetime.date(value).strftime('%d %b %Y')
            return str(val)
        else:
            return str(value)
    except Exception as e:
        logging.error(f"Failed to convert value {value}: {e}")
        return ""
   

def replace_placeholders(doc, row):
    try:

        for paragraph in doc.paragraphs:
            i = 0
            while i < len(paragraph.runs):
                text_accum = ""
                run_indices = []
                j = i
                while j < len(paragraph.runs) and len(run_indices) < 10:  # Prevent runaway loop
                    text_accum += paragraph.runs[j].text
                    run_indices.append(j)
                    for key in headers:
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in text_accum:
                            replaced = text_accum.replace(placeholder, format_value(row[key]))
                            # Clear and replace in first run
                            paragraph.runs[run_indices[0]].text = replaced
                            # Clear the remaining runs
                            for k in run_indices[1:]:
                                paragraph.runs[k].text = ""
                            i = run_indices[-1] + 1  # Move index past replaced runs
                            break
                    else:
                        j += 1
                        continue
                    break
                else:
                    i += 1
    except Exception as e:
        logging.error(f"❌ Error occured in replacing data in MS-Word template")

def replace_placeholders_in_tables(doc, row):
    try:

        for table in doc.tables:
            for tbl_row in table.rows:
                for cell in tbl_row.cells:
                    for paragraph in cell.paragraphs:
                        i = 0
                        while i < len(paragraph.runs):
                            text_accum = ""
                            run_indices = []
                            j = i
                            while j < len(paragraph.runs) and len(run_indices) < 10:  # safety check
                                text_accum += paragraph.runs[j].text
                                run_indices.append(j)
                                for key in headers:
                                    placeholder = f"{{{{{key}}}}}"
                                    if placeholder in text_accum:
                                        replaced = text_accum.replace(placeholder, format_value(row[key]))
                                        paragraph.runs[run_indices[0]].text = replaced
                                        for k in run_indices[1:]:
                                            paragraph.runs[k].text = ""
                                        i = run_indices[-1] + 1
                                        break
                                else:
                                    j += 1
                                    continue
                                break
                            else:
                                i += 1
    except Exception as e:
        logging.error(f"❌ Error occured in replacing data in MS-Word template")

def generate_docs(row):
    try:
        logging.info(f'Generating document for row: {row.to_dict()}')
        doc = Document(template_file)
        replace_placeholders(doc, row)
        replace_placeholders_in_tables(doc,row)

        name = row.get("Name", f"Record_{row.name}")
        name_date = f"{name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        file_name_with_date = "".join(c if c.isalnum() else "_" for c in str(name_date))
       
        receipt_num = row.get("ReceiptNo", f"Receipt{row.name}")
        name_with_receipt = f"{name}_{datetime.now().strftime('%Y%m%d')}_{receipt_num}"

        docx_path = os.path.join(output_dir_word, f"{file_name_with_date}.docx")
        doc.save(docx_path)
        logging.info(f'DOCX generated: {docx_path}')

        # if isPdf.lower() == 'y':
        pdf_path = os.path.join(output_dir_pdf, f"{file_name_with_date}.pdf")
        try:
            convert(docx_path, pdf_path)
            logging.info(f'PDF generated: {pdf_path}')
        except AssertionError as e:
            logging.error(f"❌ AssertionError: {e}")
            logging.info("💡 Is Microsoft Word installed and activated?")
        except Exception as e:
            logging.warning(f"⚠️ PDF conversion failed for {file_name_with_date}: {e}")
        return pdf_path, name_with_receipt
    except Exception as e:
         logging.error(f"❌ Error occured in generating a Document")
def sending_email_to_receipts(row, pdf_path, name_with_receipt):
    name = row.get("Name", f"Record_{row.name}")
    # load body
    body=''
    try:
        with open(body_file_path, 'r', encoding='utf-8') as f:
            body = f.read()
        logging.info(f'Loaded email body from {body_file_path}')
    except FileNotFoundError:
        logging.error(f"Error: File not found at '{body_file_path}'")

      # Send email with attachment (customize recipient and message as needed)
    to_email = row.get('Email', None)  # Assumes 'Email' column exists
    if to_email and not pd.isna(to_email) and str(to_email).strip() and is_valid_email_syntax(to_email) :
        subject = f"Donation Receipt - {name_with_receipt}"
        body_msg = f"Dear {name}, \n \n{body}"
        try:
            send_email_with_attachment(to_email, subject, body_msg, pdf_path)
            logging.info(f"Email sent to {to_email} with attachment {pdf_path}")
        except Exception as e:
            logging.error(f"Failed to send email to {to_email}: {e}")
    else:
        logging.warning(f'Invalid email or Sender email address missing: {to_email}')

# --- Generate for All Rows ---
logging.info('Starting document generation and email sending for all rows.')
for _, row in df.iterrows():
    pdf_path, name_with_receipt = generate_docs(row)
    
    # sending email
    if initial_email:
        sending_email_to_receipts(row, pdf_path, name_with_receipt)

logging.info(f"✅ All Word and PDF files generated in: {output_dir_word}, {output_dir_pdf}")
logging.info('Script finished.')
print('Script finished.')

